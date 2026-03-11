from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


REPO_ROOT = Path(__file__).resolve().parent


def set_document_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for section in doc.sections:
        section.page_height = Inches(11.69)  # A4
        section.page_width = Inches(8.27)
        section.top_margin = Inches(0.79)
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(0.79)
        section.right_margin = Inches(0.79)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    for r in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(r):
            row_cells[i].text = val or ""


def add_steps(doc: Document, steps: list[str]) -> None:
    for s in steps:
        doc.add_paragraph(s, style="List Number")


def main() -> None:
    doc = Document()
    set_document_defaults(doc)

    # Címlap
    doc.add_heading("HotelFlow – Felhasználói dokumentáció", level=0)
    doc.add_paragraph("Dokumentum típusa: felhasználói kézikönyv (nem technikai).")
    doc.add_paragraph("Cél: bemutatni, hogyan használható a rendszer vendégként, szálloda adminisztrátorként és super adminként.")
    doc.add_paragraph("Verzió: a repóban található jelenlegi megvalósítás alapján.")
    doc.add_page_break()

    # 1. Bevezetés
    doc.add_heading("1. Bevezetés", level=1)
    doc.add_paragraph("A rendszer neve: HotelFlow – Szálloda Foglalási Rendszer.")
    doc.add_paragraph(
        "A rendszer célja, hogy a vendégek egyszerűen szálláshelyet keressenek és foglaljanak, "
        "a szálloda adminisztrátorok pedig kezeljék a foglalásokat, szobákat, szolgáltatásokat, számlázást és RFID kulcsokat."
    )
    doc.add_paragraph(
        "Mire használható: szálláshely keresés, foglalási kérelem leadása, foglalások nyomon követése, "
        "számla fogadása (PDF), bankkártyás fizetés fizetési linken keresztül, illetve check-in QR kód e-mailben (fizetés után)."
    )
    doc.add_paragraph("A dokumentáció célja: lépésről lépésre megmutatni a rendszer használatát nem technikai felhasználóknak.")

    # 2. Rendszerkövetelmények
    doc.add_heading("2. Rendszerkövetelmények", level=1)
    doc.add_paragraph("A rendszer webes alkalmazás, modern böngészőből használható.")
    doc.add_paragraph("Támogatott böngészők (ajánlott):")
    doc.add_paragraph("Google Chrome (friss verzió)", style="List Bullet")
    doc.add_paragraph("Microsoft Edge (friss verzió)", style="List Bullet")
    doc.add_paragraph("Mozilla Firefox (friss verzió)", style="List Bullet")
    doc.add_paragraph("Safari (macOS/iOS – friss verzió)", style="List Bullet")
    doc.add_paragraph("Internet kapcsolat: stabil, folyamatos kapcsolat javasolt (foglalás és fizetés során kiemelten).")
    doc.add_paragraph("Ajánlott eszközök:")
    doc.add_paragraph("PC/laptop (admin felülethez különösen kényelmes)", style="List Bullet")
    doc.add_paragraph("Mobiltelefon (vendég oldali böngészés/fizetés támogatott, reszponzív felületekkel)", style="List Bullet")

    # 3. Bejelentkezés
    doc.add_heading("3. Bejelentkezés a rendszerbe", level=1)
    doc.add_heading("3.1 Bejelentkezés lépései", level=2)
    add_steps(
        doc,
        [
            "Nyissa meg a bejelentkezési oldalt (Bejelentkezés).",
            "Adja meg az e-mail címét.",
            "Adja meg a jelszavát.",
            "Kattintson a „BEJELENTKEZÉS” gombra.",
        ],
    )
    doc.add_heading("3.2 Sikeres bejelentkezés után mi történik?", level=2)
    doc.add_paragraph("A rendszer a szerepkör alapján automatikusan a megfelelő felületre irányít:")
    add_table(
        doc,
        headers=["Szerepkör", "Hova irányít a rendszer?"],
        rows=[
            ["Vendég (user)", "„Foglalásaim” oldal"],
            ["Szálloda admin (hotel)", "Admin → Foglalások"],
            ["Super admin (super_admin)", "Super admin → Dashboard"],
        ],
    )
    doc.add_heading("3.3 Hibás adatok / tipikus problémák", level=2)
    doc.add_paragraph("Ha hibás e-mail/jelszó párost ad meg, a rendszer hibaüzenetet jelenít meg.")
    doc.add_paragraph(
        "Ha a regisztráció után még nem erősítette meg az e-mail címét, a rendszer jelezni fogja, "
        "és felajánlhatja a megerősítő e-mail újraküldését."
    )
    doc.add_heading("3.4 Kétfaktoros hitelesítés (2FA)", level=2)
    doc.add_paragraph(
        "Bizonyos fiókoknál a rendszer kérhet 2FA kódot. Ilyenkor az autentikátor alkalmazás (pl. Google Authenticator/Authy) "
        "által generált 6 számjegyű kódot kell megadni."
    )
    doc.add_paragraph("Ha elvész a telefon, a rendszer 2FA helyreállítást is kínál e-mailes linken keresztül.")

    # 4. A rendszer felépítése
    doc.add_heading("4. A rendszer felépítése", level=1)
    doc.add_heading("4.1 Fő felületek (szerepkörök szerint)", level=2)
    doc.add_paragraph("Vendég (user) felületek:", style=None)
    doc.add_paragraph("Keresés (szállások keresése)", style="List Bullet")
    doc.add_paragraph("Szálloda részletek (csomagok, szobák, szolgáltatások)", style="List Bullet")
    doc.add_paragraph("Foglalásaim (foglalások és vendégadatok kezelése, számla megtekintése)", style="List Bullet")
    doc.add_paragraph("Fizetés (bankkártyás fizetési oldal a számla linken keresztül)", style="List Bullet")
    doc.add_paragraph("Profil (adatmódosítás, 2FA, jelszó csere, fiók törlés)", style="List Bullet")

    doc.add_paragraph("Szálloda admin (hotel) felületek:", style=None)
    doc.add_paragraph("Dashboard (statisztikák, legutóbbi tevékenységek, gyors műveletek)", style="List Bullet")
    doc.add_paragraph("Foglalások (foglalások kezelése, számlázás és fizetés státuszok)", style="List Bullet")
    doc.add_paragraph("Szobák / Szolgáltatások / Címkék / RFID kulcsok (kezelő oldalak)", style="List Bullet")
    doc.add_paragraph("Cégadatok / Számlázási információk (adószám, bankszámlaszám, EU adószám)", style="List Bullet")
    doc.add_paragraph("Profilom (fiókadatok, 2FA, jelszó, fiók törlés)", style="List Bullet")

    doc.add_paragraph("Super admin felületek:", style=None)
    doc.add_paragraph("Dashboard (összesített statisztikák, legutóbbi foglalások)", style="List Bullet")
    doc.add_paragraph("Felhasználók / Szállodák / Szobák / Szolgáltatások / Foglalások / Számlák / RFID / Eszközök", style="List Bullet")

    # 5. Fő funkciók
    doc.add_heading("5. A rendszer fő funkciói", level=1)

    doc.add_heading("5.1 Regisztráció (Vendég vagy Szálloda admin)", level=2)
    doc.add_paragraph("Mire használható: új fiók létrehozása a rendszerben.")
    doc.add_paragraph("Hogyan érhető el: „Regisztráció” menüpont.")
    doc.add_heading("Lépések", level=3)
    add_steps(
        doc,
        [
            "Válassza ki a szerepkört: „Vendég” vagy „Szálloda Admin”.",
            "Töltse ki a regisztrációs űrlapot (teljes név, e-mail, jelszó).",
            "Szálloda admin esetén adja meg a szálloda adatait (név, helyszín, típus, csillagok – opcionális).",
            "Kattintson a „REGISZTRÁLÁS” gombra.",
            "Ellenőrizze az e-mail fiókját, és kattintson a megerősítő linkre (e-mail cím megerősítése).",
        ],
    )

    doc.add_heading("5.2 Bejelentkezés + 2FA", level=2)
    doc.add_paragraph("Mire használható: hozzáférés a személyes/admin felületekhez.")
    doc.add_heading("Lépések", level=3)
    add_steps(
        doc,
        [
            "Adja meg az e-mail címet és jelszót.",
            "Ha a rendszer 2FA kódot kér, nyissa meg az autentikátor alkalmazást.",
            "Írja be a 6 számjegyű kódot, majd erősítse meg.",
        ],
    )

    doc.add_heading("5.3 Szállás keresése (Vendég)", level=2)
    doc.add_paragraph("Mire használható: szállodák keresése város, dátum és vendégszám alapján.")
    doc.add_heading("Lépések", level=3)
    add_steps(
        doc,
        [
            "Nyissa meg a „Keresés” oldalt.",
            "Adja meg: „Hová utazik?” (város/helyszín).",
            "Válassza ki az érkezés és távozás dátumát.",
            "Adja meg a vendégek számát (minimum 1).",
            "Kattintson a „Keresés” gombra.",
            "A találatoknál kattintson a kívánt szállodára a részletek megtekintéséhez.",
        ],
    )
    doc.add_paragraph("Tipp: a rendszer ajánlott szállásokat is megjeleníthet a keresés megkezdése előtt.")

    doc.add_heading("5.4 Csomag kiválasztása és foglalási kérelem leadása (Vendég)", level=2)
    doc.add_paragraph("Mire használható: kiválasztott szállodában csomag/szoba(ák) lefoglalása, extra szolgáltatásokkal.")
    doc.add_heading("Lépések", level=3)
    add_steps(
        doc,
        [
            "A szálloda oldalon válasszon csomagot (a csomagkártyára kattintva).",
            "Szükség esetén válasszon további szolgáltatásokat (jelölőnégyzetek).",
            "Kattintson a „Foglalás megerősítése” gombra.",
            "A felugró ablakban válassza ki a fizetési módot (banki átutalás vagy bankkártya).",
            "Adja meg a számlázási adatokat (magánszemély/cég).",
            "Erősítse meg a foglalási kérelmet.",
        ],
    )
    doc.add_paragraph(
        "Mi történik ezután: a foglalás „Függőben” státuszba kerül. A szálloda adminisztrátora elbírálja. "
        "Megerősítés után a számla e-mailben érkezik (PDF melléklettel)."
    )

    doc.add_heading("5.5 Számla és fizetés folyamata", level=2)
    doc.add_heading("5.5.1 Bankkártyás fizetés", level=3)
    doc.add_paragraph(
        "A számla e-mail tartalmazhat fizetési linket. A link megnyitásával elérhető a „Bankkártyás fizetés” oldal."
    )
    add_steps(
        doc,
        [
            "Nyissa meg az e-mailben kapott fizetési linket.",
            "Ellenőrizze a számla részleteit (sorszámla, összeg, határidő).",
            "Töltse ki a kártyaadatokat: kártyaszám, lejárat, CVV, kártyatulajdonos neve.",
            "Kattintson a „Fizetés” gombra.",
            "Siker esetén a rendszer visszajelzést ad („Fizetés sikeres!”).",
        ],
    )

    doc.add_heading("5.5.2 Banki átutalás", level=3)
    doc.add_paragraph(
        "Banki átutalás esetén a vendég e-mailben kapja meg a számlát (PDF), és a fizetést a számlán szereplő "
        "bankszámlaszámra kell teljesíteni. A fizetés beérkezését a szálloda adminisztrátora igazolja vissza."
    )
    doc.add_paragraph(
        "Fontos: a check-in QR kód e-mail kiküldése a fizetés megerősítése után történik."
    )

    doc.add_heading("5.6 Foglalásaim kezelése (Vendég)", level=2)
    doc.add_paragraph("Mire használható: saját foglalások áttekintése, szűrése, vendégek kezelése, számla megtekintése.")
    doc.add_heading("Lépések – Foglalások áttekintése", level=3)
    add_steps(
        doc,
        [
            "Nyissa meg a „Foglalásaim” oldalt.",
            "Használja a szűrőgombokat (Összes / Függőben / Megerősítve / Törölve / Befejezve).",
            "Válasszon nézetet: „Kártya” vagy „Táblázat”.",
            "Nyissa meg a foglalás műveleteit a „Műveletek” gombbal (ha elérhető).",
        ],
    )
    doc.add_heading("Lépések – Vendégek hozzáadása egy foglaláshoz", level=3)
    add_steps(
        doc,
        [
            "Keresse meg a foglalást a listában.",
            "Kattintson a „+ Vendég hozzáadása” gombra.",
            "Adja meg a vendég adatait (név, azonosító okmány/ID, születési dátum).",
            "Mentse a vendéget.",
        ],
    )
    doc.add_paragraph("A rendszer jelzi, ha elérte a maximális vendégkapacitást (szobák kapacitása alapján).")

    doc.add_heading("5.7 Profil kezelése", level=2)
    doc.add_paragraph("Mire használható: név/e-mail módosítás, jelszó csere, 2FA kezelése, fiók törlése.")
    doc.add_heading("Lépések – 2FA bekapcsolása", level=3)
    add_steps(
        doc,
        [
            "Nyissa meg a „Profilom” oldalt.",
            "Kapcsolja be a „Kétfaktoros hitelesítés” opciót.",
            "Kérjen QR kódot, olvassa be az autentikátor alkalmazással.",
            "Adja meg a 6 számjegyű kódot, majd erősítse meg.",
        ],
    )
    doc.add_paragraph("Szálloda adminoknál a 2FA kötelező lehet az admin felület használatához.")

    doc.add_heading("5.8 Szálloda admin: foglalások kezelése (áttekintés)", level=2)
    doc.add_paragraph("Mire használható: foglalások elfogadása/elutasítása, számlázási folyamatok kezelése, fizetés megerősítése átutalásnál.")
    doc.add_heading("Tipikus lépések", level=3)
    add_steps(
        doc,
        [
            "Lépjen be szálloda admin fiókkal.",
            "Nyissa meg az Admin → Foglalások oldalt.",
            "Szűrje a foglalásokat státusz alapján (pl. Függőben / Megerősítve / Fizetésre vár).",
            "Nyissa meg egy foglalás részleteit/műveleteit, és végezze el a szükséges admin lépést (elfogadás/elutasítás, számla előnézet/jóváhagyás/küldés).",
            "Átutalás esetén a beérkezett fizetést a rendszerben adminisztrátorként erősítse meg (ezután megy ki a check-in QR e-mail).",
        ],
    )

    doc.add_heading("5.9 Szálloda admin: számlázási/cégadatok megadása", level=2)
    doc.add_paragraph("Mire használható: a számlák kiállításához szükséges adatok rögzítése.")
    doc.add_heading("Lépések", level=3)
    add_steps(
        doc,
        [
            "Nyissa meg az Admin → „Cégadatok / Számlázási információk” oldalt.",
            "Válassza ki a szállodát (ha több szállodája van).",
            "Töltse ki a kötelező mezőket: Adószám, Bankszámlaszám, Közösségi adószám.",
            "Kattintson a „Mentés” gombra.",
        ],
    )

    doc.add_heading("5.10 Super admin: kezelőfunkciók (áttekintés)", level=2)
    doc.add_paragraph(
        "A super admin felület célja a rendszer egészének felügyelete. Itt kezelhetők többek között a felhasználók, "
        "szállodák, foglalások és egyéb törzsadatok."
    )
    doc.add_paragraph("Példa: felhasználó hozzáadása (Super admin → Felhasználók):", style=None)
    add_steps(
        doc,
        [
            "Nyissa meg a „Felhasználók kezelése” oldalt.",
            "Kattintson a „Felhasználó hozzáadása” gombra.",
            "Adja meg a nevet, e-mail címet, szerepkört, és (új felhasználónál) jelszót.",
            "Mentse a felhasználót.",
        ],
    )

    # 6. Űrlapok használata
    doc.add_heading("6. Űrlapok használata", level=1)
    doc.add_paragraph("Általános szabályok:")
    doc.add_paragraph("A * jelöli a kötelező mezőket (ahol a felület ezt kiírja).", style="List Bullet")
    doc.add_paragraph("E-mail mezőbe valós e-mail formátum szükséges (pl. pelda@email.com).", style="List Bullet")
    doc.add_paragraph("Jelszó: minimum 8 karakter.", style="List Bullet")
    doc.add_paragraph("Dátummezők: érkezés ≤ távozás, és az érkezés nem lehet múltbeli.", style="List Bullet")
    doc.add_paragraph("Szám mezők (vendégszám): minimum 1.", style="List Bullet")

    doc.add_heading("6.1 Gyakori mezők és jelentésük", level=2)
    add_table(
        doc,
        headers=["Mező", "Mit jelent?", "Hol található?"],
        rows=[
            ["E-mail cím", "Bejelentkezéshez és értesítésekhez használt cím", "Login / Regisztráció / Profil"],
            ["Jelszó", "Fiók védelme", "Login / Regisztráció / Profil / Jelszó visszaállítás"],
            ["Érkezés / Távozás dátuma", "A foglalás időtartama", "Keresés / Szálloda oldal"],
            ["Vendégek száma", "Hány főre keres szállást", "Keresés / Szálloda oldal"],
            ["Számlázási adatok", "A számlán szereplő adatok (magánszemély/cég)", "Foglalás megerősítése (modal)"],
        ],
    )

    doc.add_heading("6.2 Hibakezelés űrlapoknál", level=2)
    doc.add_paragraph("Ha egy mező hibás vagy hiányzik:")
    doc.add_paragraph("A rendszer hibaüzenetet jelenít meg a mező mellett vagy az űrlap tetején.", style="List Bullet")
    doc.add_paragraph("Javítsa a megadott adatot a hibaüzenet alapján, majd próbálja újra.", style="List Bullet")

    # 7. Hibaüzenetek és problémák
    doc.add_heading("7. Hibaüzenetek és problémák", level=1)
    doc.add_heading("7.1 Hibás bejelentkezés", level=2)
    doc.add_paragraph("Tünet: „Érvénytelen bejelentkezési adatok” vagy hasonló üzenet.")
    doc.add_paragraph("Megoldás:")
    add_steps(
        doc,
        [
            "Ellenőrizze, hogy helyes-e az e-mail cím.",
            "Ellenőrizze, hogy a Caps Lock nincs-e bekapcsolva.",
            "Ha nem emlékszik a jelszóra, használja az „Elfelejtetted a jelszavad?” funkciót.",
        ],
    )

    doc.add_heading("7.2 E-mail nincs megerősítve", level=2)
    doc.add_paragraph("Tünet: a rendszer kéri az e-mail megerősítését.")
    doc.add_paragraph("Megoldás:")
    add_steps(
        doc,
        [
            "Keresse meg a megerősítő e-mailt a postaládájában (Spam/Promóciók mappát is).",
            "Kattintson a megerősítő linkre.",
            "Ha nem érkezett meg, a bejelentkezési oldalon kérje az e-mail újraküldését.",
        ],
    )

    doc.add_heading("7.3 Hiányzó adatok / hibás formátum", level=2)
    doc.add_paragraph("Példák:")
    doc.add_paragraph("Kötelező mező üresen maradt.", style="List Bullet")
    doc.add_paragraph("Dátumok felcserélve vagy múltbeli érkezés.", style="List Bullet")
    doc.add_paragraph("Jelszó túl rövid (min. 8 karakter).", style="List Bullet")
    doc.add_paragraph("Megoldás: javítsa a mező(ke)t, majd mentse/folytassa újra.")

    doc.add_heading("7.4 Fizetési oldal hiba", level=2)
    doc.add_paragraph("Tünet: a fizetési link megnyitásakor „Nem sikerült betölteni a számla adatait” vagy más hiba.")
    doc.add_paragraph("Megoldás:")
    add_steps(
        doc,
        [
            "Ellenőrizze, hogy a link teljes és nem sérült (ne hagyjon le karaktereket).",
            "Próbálja meg másik böngészőben megnyitni.",
            "Ha továbbra sem működik, kérjen új linket a szállodától (számla újraküldése).",
        ],
    )

    # 8. Kijelentkezés
    doc.add_heading("8. Kijelentkezés", level=1)
    doc.add_paragraph("A kijelentkezés a fiók biztonsága miatt különösen fontos nyilvános vagy megosztott eszközön.")
    doc.add_heading("Lépések", level=2)
    add_steps(
        doc,
        [
            "Keresse meg a navigációban a „Kijelentkezés” opciót (ahol elérhető).",
            "Kattintson a kijelentkezésre.",
            "Zárja be a böngészőt, ha megosztott gépet használ.",
        ],
    )

    # 9. GYIK
    doc.add_heading("9. Gyakori kérdések (GYIK)", level=1)
    doc.add_heading("9.1 Nem kapom meg a megerősítő e-mailt. Mit tegyek?", level=2)
    doc.add_paragraph("Ellenőrizze a Spam/Promóciók mappát. A bejelentkezési oldalon kérje az e-mail újraküldését.")
    doc.add_heading("9.2 Mikor kapok check-in QR kódot?", level=2)
    doc.add_paragraph("A check-in QR kód e-mailt a rendszer a fizetés megerősítése után küldi ki.")
    doc.add_heading("9.3 Banki átutalásnál miért nem jön azonnal a QR kód?", level=2)
    doc.add_paragraph("Átutalásnál a fizetést a szálloda adminisztrátora igazolja vissza, utána érkezik a QR kód.")
    doc.add_heading("9.4 Elfelejtettem a jelszavam. Hogyan tudom visszaállítani?", level=2)
    doc.add_paragraph("A bejelentkezés oldalon válassza az „Elfelejtetted a jelszavad?” lehetőséget, és kövesse az e-mailben kapott lépéseket.")

    # 10. Összefoglalás
    doc.add_heading("10. Összefoglalás", level=1)
    doc.add_paragraph(
        "A HotelFlow célja, hogy egyszerű és átlátható módon támogassa a szállás keresést, a foglalást, "
        "valamint a szállodai adminisztrációt. A felhasználók gyorsan elérik a számukra fontos funkciókat, "
        "a foglalás állapota követhető, a számla e-mailben érkezik, a kártyás fizetés pedig a fizetési linken keresztül elvégezhető."
    )

    out_path = REPO_ROOT / "HotelFlow_Felhasznaloi_Dokumentacio.docx"
    doc.save(str(out_path))
    print(f"OK: {out_path}")


if __name__ == "__main__":
    main()

