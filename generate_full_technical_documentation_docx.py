from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List, Optional, Tuple

from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.shared import OxmlElement, qn


REPO_ROOT = Path(__file__).resolve().parent


def read_text(rel_path: str) -> str:
    p = REPO_ROOT / rel_path
    return p.read_text(encoding="utf-8", errors="replace")


def safe_trim(s: str, max_len: int = 500) -> str:
    s = (s or "").strip()
    if len(s) <= max_len:
        return s
    return s[: max_len - 3].rstrip() + "..."


def set_document_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Set A4 + margins (Word default-ish)
    for section in doc.sections:
        section.page_height = Inches(11.69)  # A4 height
        section.page_width = Inches(8.27)  # A4 width
        section.top_margin = Inches(0.79)
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(0.79)
        section.right_margin = Inches(0.79)


def add_code_block(doc: Document, code: str) -> None:
    """
    Word doesn't have a built-in "code block" style by default.
    We simulate it with a monospace font + light shading.
    """
    p = doc.add_paragraph()
    run = p.add_run(code.rstrip() + "\n")
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)

    # Shading
    p_format = p.paragraph_format
    p_format.space_before = Pt(4)
    p_format.space_after = Pt(4)

    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), "F3F4F6")  # light gray
    p._p.get_or_add_pPr().append(shading_elm)


@dataclass
class ApiRoute:
    method: str
    path: str
    controller: str
    action: str
    middlewares: str


def parse_api_routes(api_php: str) -> List[ApiRoute]:
    """
    Best-effort parser for typical Laravel Route::<method> definitions.
    We don't attempt to resolve nested groups fully; we capture what's visible on the same line.
    """
    routes: List[ApiRoute] = []
    for line in api_php.splitlines():
        line = line.strip()
        if not line.startswith("Route::"):
            continue

        # Example:
        # Route::post('/auth/login', [AuthController::class, 'login']);
        m = re.search(
            r"Route::(get|post|put|patch|delete)\(\s*'([^']+)'\s*,\s*\[([A-Za-z0-9_\\\\]+)::class\s*,\s*'([^']+)'\]\s*\)(.*);",
            line,
            flags=re.IGNORECASE,
        )
        if not m:
            continue

        method, path, controller, action, tail = m.groups()
        middlewares = ""
        mw = re.search(r"->middleware\(([^)]+)\)", tail)
        if mw:
            middlewares = safe_trim(mw.group(1), 160)

        routes.append(
            ApiRoute(
                method=method.upper(),
                path=path,
                controller=controller.split("\\")[-1],
                action=action,
                middlewares=middlewares,
            )
        )
    return routes


@dataclass
class DbmlTable:
    name: str
    note: str
    columns: List[Tuple[str, str]]  # (name, type/attrs)


def parse_dbml_tables(dbml: str) -> List[DbmlTable]:
    tables: List[DbmlTable] = []

    # Table <name> { ... }
    table_re = re.compile(r"Table\s+([A-Za-z0-9_]+)\s*\{([\s\S]*?)\n\}", re.MULTILINE)
    note_re = re.compile(r"Note:\s*'([^']*)'")

    for m in table_re.finditer(dbml):
        name = m.group(1)
        body = m.group(2)

        note_m = note_re.search(body)
        note = note_m.group(1).strip() if note_m else ""

        columns: List[Tuple[str, str]] = []
        for raw in body.splitlines():
            line = raw.strip()
            if not line or line.startswith("//") or line.startswith("indexes") or line.startswith("Note:"):
                continue
            # Column line: <col> <type> [attrs...]
            col_m = re.match(r"([A-Za-z0-9_]+)\s+(.+)$", line)
            if not col_m:
                continue
            col_name, rest = col_m.groups()
            # Skip composite PK lines in indexes section already filtered above
            columns.append((col_name, safe_trim(rest, 180)))

        tables.append(DbmlTable(name=name, note=note, columns=columns))

    return tables


def add_table(doc: Document, headers: List[str], rows: List[List[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h

    for r in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(r):
            row_cells[i].text = val or ""


def main() -> None:
    api_php = read_text("backend/hotelflow_server/routes/api.php")
    dbml = read_text("database_documentation.dbml")
    frontend_readme = read_text("frontend/README.md")
    backend_composer = read_text("backend/hotelflow_server/composer.json")
    frontend_pkg = read_text("frontend/package.json")

    routes = parse_api_routes(api_php)
    tables = parse_dbml_tables(dbml)

    doc = Document()
    set_document_defaults(doc)

    # Cover / címoldal
    doc.add_heading("HotelFlow – Teljes technikai dokumentáció (vizsgaremek)", level=0)
    doc.add_paragraph("Dokumentum típusa: fejlesztői/technikai dokumentáció (backend + frontend + adatbázis).")
    doc.add_paragraph("Forrás: a projekt repójában található megvalósítás (Laravel backend, Vue 3 frontend).")
    doc.add_paragraph("Készítés: automatikus generálás + kódbázis-elemzés alapján.")
    doc.add_page_break()

    # 1. Bevezetés
    doc.add_heading("1. Bevezetés", level=1)
    doc.add_paragraph("A projekt neve: HotelFlow – Szálloda foglalási rendszer.")
    doc.add_paragraph(
        "A rendszer célja egy olyan webes alkalmazás biztosítása, amelyben vendégek szálláshelyeket kereshetnek, "
        "foglalási kérést adhatnak le, a szálloda adminisztrátora kezelheti a foglalásokat, és a rendszer számlát "
        "generál, illetve e-mail értesítéseket küld a foglalási folyamat kulcslépéseinél."
    )
    doc.add_paragraph(
        "Megoldandó probléma: a foglalási és adminisztrációs folyamatok digitalizálása, "
        "számlázási dokumentum előállítása (PDF), valamint a fizetés/igazolás és check-in folyamat támogatása."
    )
    doc.add_paragraph("Célcsoportok:")
    doc.add_paragraph("Vendégek (user)", style="List Bullet")
    doc.add_paragraph("Szálloda adminisztrátorok (hotel)", style="List Bullet")
    doc.add_paragraph("Szuper adminisztrátor (super_admin)", style="List Bullet")
    doc.add_paragraph(
        "Rövid áttekintés: a frontend Vue 3 alkalmazás, amely REST API-n keresztül kommunikál a Laravel backenddel. "
        "Az adatokat MySQL adatbázis tárolja (docker-compose konfiguráció szerint)."
    )

    # 2. Rendszer architektúra
    doc.add_heading("2. Rendszer architektúra", level=1)
    doc.add_heading("2.1 Általános felépítés", level=2)
    doc.add_paragraph(
        "A rendszer kliens–szerver architektúrára épül: a böngészőben futó kliens (frontend) HTTP-n keresztül "
        "hívja a szerveroldali REST API-t (backend). A backend hitelesítést és jogosultságkezelést végez, "
        "üzleti logikát futtat, majd adatbázisban perzisztál."
    )
    doc.add_heading("2.2 Fő komponensek", level=2)
    doc.add_paragraph("Frontend (Vue 3 + Vite): felhasználói felület, űrlapok, navigáció, API hívások.")
    doc.add_paragraph("Backend (Laravel 12): API, hitelesítés (Sanctum), számla PDF generálás, e-mail küldés.")
    doc.add_paragraph("Adatbázis (MySQL): üzleti adatok (felhasználók, szállodák, szobák, foglalások, számlák, RFID).")
    doc.add_heading("2.2.1 Kliens–szerver működés (folyamat-ábra szövegesen)", level=3)
    doc.add_paragraph("1) A felhasználó a frontend felületen műveletet indít (pl. foglalás).")
    doc.add_paragraph("2) A frontend a `services/*` rétegen keresztül HTTP kérést küld a backend `/api/...` végpontjára.")
    doc.add_paragraph("3) A backend validál, auth-ot és role-t ellenőriz (`auth:sanctum`, `role:*`).")
    doc.add_paragraph("4) Üzleti logika fut (árkalkuláció, kapcsolótáblák mentése, PDF generálás, e-mail küldés).")
    doc.add_paragraph("5) A backend JSON választ ad, a frontend pedig frissíti a UI-t.")
    doc.add_heading("2.3 Frontend–backend kapcsolat", level=2)
    doc.add_paragraph(
        "A frontend az `axios` könyvtárat használja. A `frontend/src/services/api.js` interceptor automatikusan "
        "csatolja a Bearer tokent az `Authorization` headerben, ha a felhasználó be van jelentkezve."
    )

    # 3. Felhasznált technológiák
    doc.add_heading("3. Felhasznált technológiák", level=1)
    doc.add_heading("3.1 Backend technológiák", level=2)
    doc.add_paragraph("Nyelv: PHP 8.2+")
    doc.add_paragraph("Framework: Laravel 12")
    doc.add_paragraph("Adatbázis: MySQL 8 (fejlesztésben docker-compose)")
    doc.add_paragraph("ORM: Eloquent ORM (modellek: `app/Models/*`)")
    doc.add_paragraph("Hitelesítés: Laravel Sanctum (token alapú)")
    doc.add_paragraph("PDF: dompdf (számla PDF)")
    doc.add_paragraph("QR: endroid/qr-code + simple-qrcode (különböző QR felhasználások)")
    doc.add_heading("3.2 Frontend technológiák", level=2)
    doc.add_paragraph("Framework: Vue 3")
    doc.add_paragraph("Build tool: Vite")
    doc.add_paragraph("Routing: Vue Router")
    doc.add_paragraph("State: egyszerű reactive store (`src/stores/auth.js`)")
    doc.add_paragraph("API kommunikáció: Axios (service layer: `src/services/*`)")
    doc.add_heading("3.3 Verziók / függőségek (a kódból)", level=2)
    doc.add_paragraph("A backend fő csomagjai a `backend/hotelflow_server/composer.json` alapján:")
    try:
        composer_lines = "\n".join(backend_composer.splitlines()[0:60])
        add_code_block(doc, composer_lines)
    except Exception:
        pass
    doc.add_paragraph("A frontend fő függőségei a `frontend/package.json` alapján:")
    try:
        pkg_lines = "\n".join(frontend_pkg.splitlines()[0:80])
        add_code_block(doc, pkg_lines)
    except Exception:
        pass

    # 4. Backend dokumentáció
    doc.add_heading("4. Backend dokumentáció", level=1)
    doc.add_heading("4.1 Backend projekt struktúra", level=2)
    doc.add_paragraph("Gyökér: `backend/hotelflow_server` (Laravel alkalmazás).")
    doc.add_paragraph("Fontos mappák:", style=None)
    for item in [
        "app/Http/Controllers – API végpontok megvalósítása (kontrollerek)",
        "app/Models – Eloquent modellek/entitások",
        "app/Mail – e-mail sablonokat összefogó Mailables osztályok",
        "resources/views – Blade nézetek (email, PDF számla)",
        "routes/api.php – API route definíciók",
        "database/migrations – adatbázis séma változásai",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("4.2 Adatbázis felépítése", level=2)
    doc.add_paragraph(
        "A projekt tartalmaz egy dbdiagram.io kompatibilis DBML állományt (`database_documentation.dbml`), "
        "amely a táblákat, mezőket és kapcsolatokat leírja. A rendszer fő táblacsoportjai:"
    )
    doc.add_paragraph("Felhasználók és hitelesítés (users, password_reset_tokens, two_factor_recovery_tokens)", style="List Bullet")
    doc.add_paragraph("Szállodák/szobák/szolgáltatások (hotels, rooms, services, képek és tag kapcsolatok)", style="List Bullet")
    doc.add_paragraph("Foglalások (bookings + kapcsolótáblák)", style="List Bullet")
    doc.add_paragraph("Számlázás és fizetés (invoices, booking_payments, booking_invoice_details)", style="List Bullet")
    doc.add_paragraph("RFID és eszközök (rfid_keys, assignments, connections, devices…)", style="List Bullet")

    # DB táblák összefoglaló táblázat (teljes lista)
    doc.add_heading("4.2.1 Táblák összefoglalója (teljes lista)", level=3)
    add_table(
        doc,
        headers=["Tábla", "Leírás (Note)"],
        rows=[[t.name, t.note] for t in tables],
    )
    doc.add_paragraph("Megjegyzés: a táblák mezőszintű részletezését a 6. fejezet adja meg (részletes mintával).")

    doc.add_heading("4.3 Modellek / entitások", level=2)
    doc.add_paragraph(
        "A backend Eloquent modelleket használ. Példák a fő entitásokra és kapcsolataikra (a kódból):"
    )
    doc.add_paragraph("- User: Sanctum tokenek, szerepkör (user/hotel/super_admin), e-mail verifikáció, 2FA mezők.")
    doc.add_paragraph("- Booking: user–hotel kapcsolat, rooms/services many-to-many, guests hasMany, invoice/payment hasOne.")
    doc.add_paragraph("- Invoice: booking-hoz tartozó számla, státusz (draft/approved/sent), PDF útvonal, payment_token.")

    # 4.4 API végpontok
    doc.add_heading("4.4 API végpontok", level=2)
    doc.add_paragraph("Az API végpontok a `routes/api.php` fájlban vannak. Az alábbiakban csoportosítva dokumentáljuk őket.")

    def add_routes_group(title: str, path_prefix: str) -> None:
        doc.add_heading(title, level=3)
        grp = [r for r in routes if r.path.startswith(path_prefix)]
        if not grp:
            doc.add_paragraph("Nincs találat a kódban ehhez a prefixhez.")
            return
        add_table(
            doc,
            headers=["HTTP", "Végpont", "Handler", "Middleware"],
            rows=[[r.method, r.path, f"{r.controller}@{r.action}", r.middlewares] for r in grp[:120]],
        )

    add_routes_group("4.4.1 Auth végpontok", "/auth/")
    add_routes_group("4.4.2 Foglalás (bookings) végpontok", "/bookings")
    add_routes_group("4.4.3 Számla / fizetés (invoices) végpontok", "/invoices")
    add_routes_group("4.4.4 Szálloda (hotels) végpontok", "/hotels")
    add_routes_group("4.4.5 Szobák (rooms) végpontok", "/rooms")
    add_routes_group("4.4.6 Szolgáltatások (services) végpontok", "/services")
    add_routes_group("4.4.7 Címkék (tags) végpontok", "/tags")
    add_routes_group("4.4.8 Képek (images) végpontok", "/images")
    add_routes_group("4.4.9 Keresés (search) és ajánlások (recommendations)", "/search")
    add_routes_group("4.4.10 RFID kulcsok", "/rfid-keys")
    add_routes_group("4.4.11 Eszközök", "/devices")
    add_routes_group("4.4.12 Super admin", "/super-admin/")

    doc.add_heading("4.4.13 Példa kérés/válasz formátumok", level=3)
    doc.add_paragraph("Megjegyzés: a konkrét mezők endpointonként eltérnek; az alábbi minták a kódban megfigyelt általános JSON sémát mutatják.")
    add_code_block(
        doc,
        "POST /api/auth/login\n"
        "Request JSON: { \"email\": \"...\", \"password\": \"...\", \"two_factor_code\"?: \"123456\" }\n"
        "Response JSON (siker): { \"id\": 1, \"name\": \"...\", \"role\": \"user|hotel|super_admin\", \"token\": \"...\" }\n"
    )
    add_code_block(
        doc,
        "POST /api/bookings\n"
        "Request JSON (részlet): { userId, hotelId, startDate, endDate, rooms: [{id, guests}], services?: [], payment_method?: \"card|bank_transfer\", invoice_details?: {...} }\n"
        "Response JSON: { \"bookingId\": 123, \"totalPrice\": 999 }\n"
    )

    doc.add_heading("4.5 Hitelesítés és jogosultságkezelés", level=2)
    doc.add_paragraph("Bejelentkezés / token:")
    doc.add_paragraph(
        "A login végpont a felhasználó jelszavát ellenőrzi, majd Sanctum tokent generál. "
        "A token a frontendben `localStorage`-be kerül, és az axios interceptor automatikusan csatolja."
    )
    doc.add_paragraph("E-mail megerősítés:")
    doc.add_paragraph(
        "Regisztráció után a felhasználó e-mail verifikációs linket kap, és csak verifikáció után jelentkezhet be "
        "(kivétel: super_admin szerepkör)."
    )
    doc.add_paragraph("2FA:")
    doc.add_paragraph(
        "A super_admin számára kötelező a 2FA, hotel adminoknál a frontend router guard kényszerítheti a bekapcsolást "
        "(a kódban szereplő szabály szerint)."
    )
    doc.add_paragraph("Szerepkör ellenőrzés:")
    doc.add_paragraph(
        "A backend `role:*` middleware-t használ (lásd `app/Http/Middleware/RoleMiddleware.php`), "
        "amely a super_admin számára teljes hozzáférést enged, egyébként szerepkör-egyezést vár."
    )

    doc.add_heading("4.6 Backend üzleti logika", level=2)
    doc.add_paragraph("A legfontosabb üzleti folyamatok (a kódból levezetve):")
    doc.add_paragraph("Foglalási kérelem létrehozás", style="List Bullet")
    doc.add_paragraph(
        "A `BookingController@store` validálja a dátumokat, szobákat, szolgáltatásokat, kiszámolja a teljes árat, "
        "létrehozza a foglalást, rögzíti a fizetési módot és (opcionálisan) számlázási adatokat, "
        "és e-mailt küld a vendégnek, illetve a hotel adminnak."
    )
    doc.add_paragraph("Foglalás megerősítése (hotel admin)", style="List Bullet")
    doc.add_paragraph(
        "A `BookingController@updateStatus` megerősítéskor biztosítja a payment rekordot, "
        "és a fizetési módtól függően automatikusan generál/approve-ol/küld számlát (kártya: payment linkkel; "
        "átutalás: átutalásos email). A check-in QR e-mailt a rendszer csak sikeres fizetés után küldi."
    )
    doc.add_paragraph("Fizetés megerősítése (átutalás esetén manuális)", style="List Bullet")
    doc.add_paragraph(
        "A `BookingController@confirmPayment` hotel admin jogosultsággal `paid` státuszra állítja a fizetést, "
        "és ekkor küldi a `BookingConfirmationMail`-t (QR kód)."
    )
    doc.add_paragraph("Számla kezelés", style="List Bullet")
    doc.add_paragraph(
        "Az `InvoiceController` tud számlát előnézetben generálni, frissíteni (draft), jóváhagyni (approve), "
        "és küldeni. Jóváhagyáskor PDF generálás történik, és a fizetési módtól függően eltérő e-mail megy ki."
    )
    doc.add_heading("4.6.1 E-mail küldések áttekintése", level=3)
    doc.add_paragraph(
        "A backend több ponton küld e-mailt (Mailables: `app/Mail/*`). Tipikus események:"
    )
    for item in [
        "Regisztráció: e-mail verifikációs levél",
        "Jelszó visszaállítás: tokenes levél",
        "2FA recovery: egyszer használatos helyreállítási link",
        "Foglalási kérelem elküldve: vendég értesítése + hotel admin értesítése",
        "Foglalás megerősítés: számla küldés (fizetési módtól függő sablon)",
        "Fizetés megerősítése után: check-in QR kód email",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # 5. Frontend dokumentáció
    doc.add_heading("5. Frontend dokumentáció", level=1)
    doc.add_heading("5.1 Frontend projekt struktúra", level=2)
    doc.add_paragraph("Gyökér: `frontend/` (Vue 3 alkalmazás).")
    doc.add_paragraph("Fontos mappák:", style=None)
    for item in [
        "src/views – oldalak (vendég, admin, super admin)",
        "src/services – API szolgáltatások (axios wrapper)",
        "src/router – route definíciók + guardok",
        "src/stores – auth state (token + user)",
        "src/layouts – admin/super-admin layoutok (navigáció, sidebar)",
        "tests – Selenium alapú E2E tesztek",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("5.2 Komponensek", level=2)
    doc.add_paragraph(
        "A projektben a fő, újrahasznosítható UI elemek a `src/components/` alatt vannak "
        "(pl. `DataTable`, `ConfirmDialog`, `Toast`, 2FA promptok, image upload)."
    )
    doc.add_heading("5.3 Routing", level=2)
    doc.add_paragraph(
        "A Vue Router útvonalai a `src/router/index.js` fájlban vannak. "
        "Meta mezők jelzik, hogy kell-e auth, és szükséges-e szerepkör. "
        "A `beforeEach` guard kezeli a beléptetést/átirányítást és a role alapú védelmet."
    )

    doc.add_heading("5.4 Állapotkezelés", level=2)
    doc.add_paragraph(
        "Az auth állapotot a `src/stores/auth.js` kezeli egy `reactive` state-tel. "
        "Induláskor betölti a `localStorage`-ből a tokent és felhasználót. "
        "Login után elmenti őket, logoutnál törli."
    )

    doc.add_heading("5.5 API kommunikáció", level=2)
    doc.add_paragraph(
        "Az API hívások a `src/services/*.js` fájlokban vannak. "
        "Az `src/services/api.js` axios instance beállítja a baseURL-t, "
        "és interceptorral hozzáadja a Bearer tokent. 401 esetén automatikus kijelentkeztetés történik."
    )

    doc.add_heading("5.6 Űrlapok és felhasználói interakciók", level=2)
    doc.add_paragraph(
        "A validáció jelentős része backend oldali (Laravel `Request::validate`). "
        "Frontend oldalon a komponensek jellemzően egyszerű ellenőrzéseket végeznek "
        "(pl. kötelező mezők, dátumok), majd hiba esetén felhasználóbarát üzenetet jelenítenek meg."
    )
    doc.add_heading("5.7 Oldalak (views) funkcionális áttekintése", level=2)
    doc.add_paragraph("Vendég oldalak:", style=None)
    for item in [
        "/search – keresés és listázás",
        "/hotel/:id – szálloda részletek + foglalás indítása",
        "/bookings – saját foglalások",
        "/payment/:token – számla fizetés token alapján",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph("Hotel admin oldalak:", style=None)
    for item in [
        "/admin – dashboard",
        "/admin/bookings – foglalások kezelése, számla műveletek",
        "/admin/rooms, /admin/services – erőforrás kezelés",
        "/admin/rfid-keys – RFID kulcsok kezelése",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph("Super admin oldalak:", style=None)
    for item in [
        "/super-admin/dashboard – statisztikák",
        "/super-admin/users/hotels/rooms/services/bookings/invoices – CRUD jellegű admin funkciók",
        "/super-admin/devices – eszköz regisztrációk",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # 6. Adatbázis séma áttekintése
    doc.add_heading("6. Adatbázis séma áttekintése", level=1)
    doc.add_paragraph(
        "A teljes adatbázis séma a `database_documentation.dbml` fájlban található. "
        "Az alábbiakban a fontosabb entitáscsoportokat és kapcsolataikat foglaljuk össze."
    )
    doc.add_heading("6.1 Kapcsolati modell (magas szint)", level=2)
    for item in [
        "users 1—N hotels (egy felhasználó több szállodát is kezelhet, kódban tipikusan 1)",
        "hotels 1—N rooms / services",
        "users 1—N bookings; hotels 1—N bookings",
        "bookings N—N rooms (bookingsRelation)",
        "bookings N—N services (servicesRelation)",
        "bookings 1—N guests",
        "bookings 1—1 invoices; bookings 1—1 booking_payments; bookings 1—1 booking_invoice_details",
        "RFID: rfid_keys 1—N rfid_assignments; booking 1—N rfid_assignments",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("6.2 Táblák és mezők (részletes minták)", level=2)
    doc.add_paragraph(
        "Az alábbiakban több kulcstábla mezőlistáját közöljük. A teljes, összes mezőt tartalmazó definíciót "
        "a DBML fájl adja meg."
    )
    # részletesebb táblázat 12 táblára (méret kontroll mellett)
    for idx, t in enumerate(tables[:12], start=1):
        doc.add_heading(f"6.2.{idx} {t.name}", level=3)
        if t.note:
            doc.add_paragraph(t.note)
        add_table(
            doc,
            headers=["Mező", "Típus / attribútumok"],
            rows=[[c, a] for c, a in t.columns[:40]],
        )

    # 7. Biztonsági megoldások
    doc.add_heading("7. Biztonsági megoldások", level=1)
    doc.add_paragraph("Hitelesítés:")
    doc.add_paragraph("Sanctum token (Bearer) – védett végpontok `auth:sanctum` middleware-rel.", style="List Bullet")
    doc.add_paragraph("Jelszavak hash-elése (`Hash::make`, `Hash::check`).", style="List Bullet")
    doc.add_paragraph("E-mail verifikáció – bejelentkezés csak megerősített emaillel (nem super_admin).", style="List Bullet")
    doc.add_paragraph("2FA – super_admin kötelező, további szerepkörök opcionális.", style="List Bullet")
    doc.add_paragraph("Jogosultságkezelés:")
    doc.add_paragraph("Role middleware: `role:hotel`, `role:super_admin` stb.", style="List Bullet")
    doc.add_paragraph("Erőforrás tulajdonlás ellenőrzés (pl. booking user_id/ hotel owner check).", style="List Bullet")
    doc.add_paragraph("Input validáció:")
    doc.add_paragraph("Laravel `validate` a kontrollerekben (dátumok, ID-k, enumok, max hossz).", style="List Bullet")

    # 8. Tesztelés
    doc.add_heading("8. Tesztelés", level=1)
    doc.add_heading("8.1 Backend tesztek (PHPUnit)", level=2)
    doc.add_paragraph(
        "A backend tartalmaz PHPUnit konfigurációt (`phpunit.xml`). A jelenlegi tesztkészlet minimális (Example testek), "
        "de a tesztkörnyezet in-memory SQLite adatbázist használ."
    )
    doc.add_heading("8.2 Frontend tesztek (Selenium)", level=2)
    doc.add_paragraph(
        "A `frontend/tests` mappában Selenium WebDriver + ChromeDriver alapú E2E tesztek vannak. "
        "A `driver.js` egy Chrome drivert épít fel, a tesztfájlok pedig konkrét user flow-kat ellenőriznek "
        "(login, navigáció, keresés validáció, szolgáltatás hozzáadás)."
    )
    doc.add_heading("8.3 Teszt futtatás (javasolt lépések)", level=2)
    doc.add_paragraph("Backend (PHPUnit):", style=None)
    add_code_block(doc, "cd backend/hotelflow_server\nphp artisan test\n")
    doc.add_paragraph("Frontend (E2E):", style=None)
    add_code_block(doc, "cd frontend\nnpm install\n# majd a tests/* futtatása a projektben kialakított mód szerint\n")

    # 9. Telepítés
    doc.add_heading("9. Telepítés (Deployment)", level=1)
    doc.add_heading("9.1 Docker alapú futtatás", level=2)
    doc.add_paragraph(
        "A projekt gyökerében található `docker-compose.yml` a három fő komponenst indítja: backend, frontend, db. "
        "A frontend nginx-szel szolgálja ki a buildelt statikus fájlokat, a backend `php artisan serve` jellegű "
        "kiszolgálással elérhető a 8000-es porton."
    )
    doc.add_heading("9.2 Lokális fejlesztői futtatás", level=2)
    doc.add_paragraph("Backend: composer install, .env beállítás, artisan key:generate, migrate, artisan serve.")
    doc.add_paragraph("Frontend: npm install, Vite dev server.")
    doc.add_heading("9.3 Környezeti változók (összefoglaló)", level=2)
    doc.add_paragraph("Backend tipikus .env kulcsok (a configok és docker-compose alapján):")
    add_code_block(
        doc,
        "APP_NAME=HotelFlow\nAPP_ENV=local\nAPP_DEBUG=true\nAPP_URL=http://localhost:8000\n"
        "FRONTEND_URL=http://localhost:3000\n"
        "DB_CONNECTION=mysql\nDB_HOST=127.0.0.1\nDB_PORT=3306\nDB_DATABASE=laravel\nDB_USERNAME=laravel\nDB_PASSWORD=secret\n"
        "MAIL_MAILER=log\nINVOICE_CURRENCY=EUR\n",
    )

    # 10. Összegzés
    doc.add_heading("10. Összegzés", level=1)
    doc.add_paragraph(
        "A HotelFlow egy több szerepkörös, kliens–szerver alapú webalkalmazás, amely a szállodai foglalási folyamat "
        "digitalizálását célozza. A backend Laravel alapokon biztosítja az API-t, jogosultságkezelést, e-mail "
        "értesítéseket és PDF számlagenerálást, míg a frontend Vue 3 segítségével valósítja meg a felhasználói "
        "felületet és admin felületeket. A DBML dokumentáció lehetővé teszi az adatmodell egyszerű áttekintését és "
        "tanári értékeléshez szükséges részletek követését."
    )

    # Kódrészletek (minták) – rövid, tényleges fájlokból
    doc.add_page_break()
    doc.add_heading("Melléklet: fontos kódrészletek (részlet)", level=1)
    doc.add_heading("A) Frontend axios interceptor (token csatolás)", level=2)
    api_js_excerpt = "\n".join(read_text("frontend/src/services/api.js").splitlines()[:60])
    add_code_block(doc, api_js_excerpt)
    doc.add_heading("B) Backend role middleware (szerepkör ellenőrzés)", level=2)
    role_mw_excerpt = "\n".join(read_text("backend/hotelflow_server/app/Http/Middleware/RoleMiddleware.php").splitlines()[:80])
    add_code_block(doc, role_mw_excerpt)

    out_path = REPO_ROOT / "HotelFlow_Technikai_Dokumentacio.docx"
    doc.save(str(out_path))
    print(f"OK: {out_path}")


if __name__ == "__main__":
    main()

